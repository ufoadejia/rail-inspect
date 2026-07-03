"""探伤报告生成：汇总检测结果，导出 .docx。

python-docx 装不上时降级为纯文本报告。
"""
import os
from datetime import datetime
from sqlalchemy.orm import Session
from app.models.tables import Workorder


def generate_report_docx(db: Session, wid: int) -> tuple:
    """生成报告，返回 (字节流, content_type, 文件扩展名)。"""
    wo = db.get(Workorder, wid)
    if not wo:
        raise ValueError("工单不存在")

    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from io import BytesIO
        return _gen_docx(wo), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"
    except ImportError:
        return _gen_text(wo).encode("utf-8"), "text/plain", "txt"


def _gen_docx(wo: Workorder) -> bytes:
    from docx import Document
    from io import BytesIO

    doc = Document()
    h = doc.add_heading("钢轨探伤作业报告", level=0)
    h.alignment = 1  # 居中

    doc.add_heading("一、作业基本信息", level=1)
    info = doc.add_table(rows=4, cols=2, style="Light Grid Accent 1")
    rows = [
        ("工单编号", f"WO-{wo.id:06d}"),
        ("线路 / 千米标", f"{wo.line_name} / {wo.kilometer}"),
        ("作业股别", wo.rail_side),
        ("作业人员", wo.assignee or "—"),
    ]
    for i, (k, v) in enumerate(rows):
        info.cell(i, 0).text = k
        info.cell(i, 1).text = v

    doc.add_heading("二、探伤检测结果", level=1)
    if wo.detections:
        t = doc.add_table(rows=1, cols=5, style="Light Grid Accent 1")
        for j, header in enumerate(["序号", "图像类型", "缺陷类型", "判级", "说明"]):
            t.cell(0, j).text = header
        for idx, det in enumerate(wo.detections, 1):
            t.add_row()
            t.cell(idx, 0).text = str(idx)
            t.cell(idx, 1).text = det.image_type
            t.cell(idx, 2).text = det.defect_type
            t.cell(idx, 3).text = det.defect_grade
            t.cell(idx, 4).text = det.description
    else:
        doc.add_paragraph("暂无检测记录。")

    doc.add_heading("三、处置建议", level=1)
    heavy = [d for d in wo.detections if d.defect_grade == "重伤"]
    light = [d for d in wo.detections if d.defect_grade == "轻伤"]
    if heavy:
        doc.add_paragraph(f"发现重伤缺陷 {len(heavy)} 处，建议立即限速并安排换轨/加固处理。", style="List Bullet")
    if light:
        doc.add_paragraph(f"发现轻伤缺陷 {len(light)} 处，建议纳入监测计划，定期复核探伤。", style="List Bullet")
    if not heavy and not light:
        doc.add_paragraph("未发现明显伤损，按周期正常巡检。", style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph(f"报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _gen_text(wo: Workorder) -> str:
    """降级：纯文本报告。"""
    lines = [
        "钢轨探伤作业报告",
        "=" * 40,
        "",
        "一、作业基本信息",
        f"  工单编号: WO-{wo.id:06d}",
        f"  线路/千米标: {wo.line_name} / {wo.kilometer}",
        f"  作业股别: {wo.rail_side}",
        f"  作业人员: {wo.assignee or '—'}",
        "",
        "二、探伤检测结果",
    ]
    if wo.detections:
        for idx, det in enumerate(wo.detections, 1):
            lines.append(f"  {idx}. [{det.image_type}] {det.defect_type} ({det.defect_grade}) - {det.description}")
    else:
        lines.append("  暂无检测记录。")

    lines.append("")
    lines.append("三、处置建议")
    heavy = [d for d in wo.detections if d.defect_grade == "重伤"]
    light = [d for d in wo.detections if d.defect_grade == "轻伤"]
    if heavy:
        lines.append(f"  - 发现重伤缺陷 {len(heavy)} 处，建议立即限速并安排换轨/加固处理。")
    if light:
        lines.append(f"  - 发现轻伤缺陷 {len(light)} 处，建议纳入监测计划，定期复核探伤。")
    if not heavy and not light:
        lines.append("  - 未发现明显伤损，按周期正常巡检。")

    lines.append("")
    lines.append(f"报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    return "\n".join(lines)
